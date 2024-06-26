import io
from typing import List
from langchain_community.document_loaders import UnstructuredFileLoader


class DocLoader(UnstructuredFileLoader):
    """
    A custom document parser that extends the UnstructuredFileLoader class.
    """

    def parser(self) -> List:
        def doc2text(filepath):
            from docx.table import _Cell, Table
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.text.paragraph import Paragraph
            from docx import Document, ImagePart
            from PIL import Image
            from io import BytesIO
            import numpy as np
            from rapidocr_onnxruntime import RapidOCR
            ocr = RapidOCR()
            if isinstance(filepath, str):
                doc = Document(filepath)
            else:
                doc = Document(io.BytesIO(filepath))
            resp = ""

            def iter_block_items(parent):
                from docx.document import Document
                if isinstance(parent, Document):
                    parent_elm = parent.element.body
                elif isinstance(parent, _Cell):
                    parent_elm = parent._tc
                else:
                    raise ValueError("DocParser parse fail")

                for child in parent_elm.iterchildren():
                    if isinstance(child, CT_P):
                        yield Paragraph(child, parent)
                    elif isinstance(child, CT_Tbl):
                        yield Table(child, parent)

            for i, block in enumerate(iter_block_items(doc)):
                if isinstance(block, Paragraph):
                    resp += block.text.strip() + "\n"
                    # images = block._element.xpath('.//pic:pic')  # 获取所有图片
                    # for image in images:
                    #     for img_id in image.xpath('.//a:blip/@r:embed'):  # 获取图片id
                    #         part = doc.part.related_parts[img_id]  # 根据图片id获取对应的图片
                    #         if isinstance(part, ImagePart):
                    #             image = Image.open(BytesIO(part._blob))
                    #             result, _ = ocr(np.array(image))
                    #             if result:
                    #                 ocr_result = [line[1] for line in result]
                    #                 resp += "\n".join(ocr_result)
                elif isinstance(block, Table):
                    for row in block.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                resp += paragraph.text.strip() + "\n"
            return resp

        text = doc2text(self.file_path)
        return text